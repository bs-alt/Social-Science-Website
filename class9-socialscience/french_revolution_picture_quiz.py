import os
import shutil
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches


def extract_images(pdf_path, image_dir):
    os.makedirs(image_dir, exist_ok=True)
    images = []
    doc = fitz.open(pdf_path)
    for page_index in range(doc.page_count):
        page = doc.load_page(page_index)
        for img_index, img in enumerate(page.get_images(full=True), start=1):
            base_image = doc.extract_image(img[0])
            image_bytes = base_image["image"]
            ext = base_image["ext"]
            image_name = f"page{page_index + 1}_img{img_index}.{ext}"
            image_path = os.path.join(image_dir, image_name)
            with open(image_path, "wb") as f:
                f.write(image_bytes)
            images.append(image_path)
    doc.close()
    return images


def build_quiz_doc(images, output_path):
    document = Document()
    document.add_heading("French Revolution Picture Quiz", level=1)
    answer_key = []

    for idx, image_path in enumerate(images, start=1):
        document.add_picture(image_path, width=Inches(4))
        question = (
            f"Q{idx}. Which event of the French Revolution is shown in the image?"
        )
        document.add_paragraph(question)
        options = [
            "A. Storming of the Bastille",
            "B. Women's March on Versailles",
            "C. The Reign of Terror",
            "D. Formation of the National Assembly",
        ]
        for opt in options:
            document.add_paragraph(opt, style="List Bullet")
        document.add_paragraph("Student answer: ______________________")
        document.add_paragraph()
        answer_key.append("A")

    document.add_page_break()
    document.add_heading("Answer Key", level=2)
    for idx, ans in enumerate(answer_key, start=1):
        document.add_paragraph(f"Q{idx}: {ans}")

    document.save(output_path)


def main():
    base_dir = os.path.dirname(__file__)
    pdf_path = os.path.join(base_dir, "History_01_The_French_revolution.pdf")
    image_dir = os.path.join(base_dir, "images")
    output_doc = os.path.join(base_dir, "French_Revolution_Picture_Quiz.docx")

    images = extract_images(pdf_path, image_dir)
    build_quiz_doc(images, output_doc)

    shutil.rmtree(image_dir)


if __name__ == "__main__":
    main()
